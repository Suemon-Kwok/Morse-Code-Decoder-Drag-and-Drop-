import subprocess
import sys

def _install(package):
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', package],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

REQUIRED = {
    'pydub':               'pydub',
    'numpy':               'numpy',
    'docx':                'python-docx',
    'morse_audio_decoder': 'morse-audio-decoder',
}

for import_name, pip_name in REQUIRED.items():
    try:
        __import__(import_name)
    except ImportError:
        print(f'Installing {pip_name}…')
        _install(pip_name)

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
import tkinter.font as tkFont
import docx
import os
import wave
import tempfile
import threading
import numpy as np

# Try importing optional audio conversion library (needed for non-WAV formats)
try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False

# Try importing the morse_audio_decoder library (used as fallback decoder)
try:
    from morse_audio_decoder.morse import MorseCode as _LibMorseCode
    LIB_DECODER_AVAILABLE = True
except ImportError:
    LIB_DECODER_AVAILABLE = False

# ─── Morse Code Dictionary ────────────────────────────────────────────────────

MORSE_CODE_DICT = {
    'A': '.-',    'B': '-...',  'C': '-.-.',  'D': '-..',
    'E': '.',     'F': '..-.',  'G': '--.',   'H': '....',
    'I': '..',    'J': '.---',  'K': '-.-',   'L': '.-..',
    'M': '--',    'N': '-.',    'O': '---',   'P': '.--.',
    'Q': '--.-',  'R': '.-.',   'S': '...',   'T': '-',
    'U': '..-',   'V': '...-',  'W': '.--',   'X': '-..-',
    'Y': '-.--',  'Z': '--..',
    '1': '.----', '2': '..---', '3': '...--', '4': '....-',
    '5': '.....', '6': '-....', '7': '--...', '8': '---..',
    '9': '----.', '0': '-----',
    ',': '--..--', '.': '.-.-.-', '?': '..--..', '/': '-..-.',
    '-': '-....-', '(': '-.--.', ')': '-.--.-', "'": '.----.',
    '!': '-.-.--', '&': '.-...', ':': '---...', ';': '-.-.-.',
    '=': '-...-',  '+': '.-.-.', '_': '..--.-', '"': '.-..-.',
    '$': '...-..-','@': '.--.-.',
}
REVERSE_MORSE = {v: k for k, v in MORSE_CODE_DICT.items()}

# ─── Audio / Morse Helpers ────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = ['.wav', '.mp3', '.ogg', '.flac', '.aac', '.m4a', '.aiff', '.mov', '.mp4']


def load_audio_as_mono_wav(input_path: str) -> str:
    """Convert any supported file to a mono 22050 Hz WAV temp file."""
    ext = os.path.splitext(input_path)[1].lower()

    if not PYDUB_AVAILABLE:
        if ext != '.wav':
            raise ImportError(
                f"pydub is required to convert {ext} files.\n"
                "Install it with:  pip install pydub\n"
                "You may also need ffmpeg installed on your system."
            )
        # Bare WAV path — validate header
        with open(input_path, 'rb') as f:
            if f.read(4) != b'RIFF':
                raise ValueError("Not a valid WAV file (missing RIFF header).")
        with wave.open(input_path, 'rb') as wf:
            if wf.getnchannels() == 1:
                return input_path, False   # already mono, no temp file

    format_map = {
        '.mp3': 'mp3', '.ogg': 'ogg', '.flac': 'flac',
        '.aac': 'aac', '.m4a': 'm4a', '.aiff': 'aiff',
        '.mov': 'mov', '.mp4': 'mp4', '.wav': 'wav',
    }
    audio = AudioSegment.from_file(input_path, format=format_map.get(ext, ext))
    audio = audio.set_channels(1).set_frame_rate(22050).set_sample_width(2)
    tmp = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
    audio.export(tmp.name, format='wav')
    return tmp.name, True   # temp file created


def _read_wav_samples(wav_path: str):
    """Read a mono WAV and return (sample_rate, float32 numpy array -1..1)."""
    with wave.open(wav_path, 'rb') as wf:
        sr = wf.getframerate()
        n  = wf.getnframes()
        sw = wf.getsampwidth()
        raw = wf.readframes(n)

    dtype = {1: np.int8, 2: np.int16, 4: np.int32}.get(sw, np.int16)
    samples = np.frombuffer(raw, dtype=dtype).astype(np.float32)
    samples /= float(np.iinfo(dtype).max)
    return sr, samples


def _detect_tone_frequency(samples: np.ndarray, sr: int) -> tuple:
    """
    Auto-detect Morse carrier tone by averaging FFT over many short windows
    and finding the dominant peak between 300–4000 Hz.
    Returns (lo, hi, peak_freq).
    """
    win = min(int(sr * 0.05), len(samples))
    n_windows = min(30, max(1, len(samples) // win))
    freqs = np.fft.rfftfreq(win, 1.0 / sr)
    band_mask = (freqs >= 300) & (freqs <= 4000)

    avg_spec = np.zeros(len(freqs))
    step = max(1, (len(samples) - win) // n_windows)
    for i in range(n_windows):
        frame = samples[i * step: i * step + win] * np.hanning(win)
        avg_spec += np.abs(np.fft.rfft(frame))
    avg_spec /= n_windows
    avg_spec[~band_mask] = 0

    peak_freq = float(freqs[np.argmax(avg_spec)])
    lo = max(200, peak_freq - 250)
    hi = peak_freq + 250
    return lo, hi, peak_freq


def _bandpass_energy(samples: np.ndarray, sr: int,
                     lo: float = 400, hi: float = 1200,
                     window_ms: float = 20) -> np.ndarray:
    """Compute short-time energy in the given frequency band with 5ms hop."""
    win = max(1, int(sr * window_ms / 1000))
    hop = int(sr * 0.005)   # fixed 5ms hop
    n_frames = (len(samples) - win) // hop + 1

    freqs = np.fft.rfftfreq(win, d=1.0 / sr)
    band  = (freqs >= lo) & (freqs <= hi)

    energy = np.zeros(n_frames, dtype=np.float32)
    for i in range(n_frames):
        frame = samples[i * hop: i * hop + win] * np.hanning(win)
        spec  = np.abs(np.fft.rfft(frame))
        energy[i] = np.sum(spec[band] ** 2)

    return energy


def _smooth(arr: np.ndarray, k: int = 5) -> np.ndarray:
    kernel = np.ones(k) / k
    return np.convolve(arr, kernel, mode='same')


def _energy_to_binary(energy: np.ndarray) -> np.ndarray:
    """Otsu-like threshold → binary ON/OFF."""
    lo, hi = energy.min(), energy.max()
    if hi - lo < 1e-9:
        return np.zeros(len(energy), dtype=np.int8)
    best_t, best_var = lo, -1.0
    for t in np.linspace(lo, hi, 50):
        above = energy[energy >= t]
        below = energy[energy <  t]
        if len(above) == 0 or len(below) == 0:
            continue
        w0, w1 = len(below) / len(energy), len(above) / len(energy)
        var = w0 * w1 * (below.mean() - above.mean()) ** 2
        if var > best_var:
            best_var, best_t = var, t
    return (energy >= best_t).astype(np.int8)


def _run_lengths_ms(binary: np.ndarray, hop_ms: float = 5):
    """Return list of (value, duration_ms) runs."""
    if len(binary) == 0:
        return []
    runs, cur_val, cur_len = [], binary[0], 1
    for v in binary[1:]:
        if v == cur_val:
            cur_len += 1
        else:
            runs.append((int(cur_val), cur_len * hop_ms))
            cur_val, cur_len = v, 1
    runs.append((int(cur_val), cur_len * hop_ms))
    return runs


def _run_lengths(binary):
    return _run_lengths_ms(binary, hop_ms=1)


def _auto_classify_and_decode(runs_ms):
    """
    Given (value, ms) runs, auto-measure dot unit then decode.
    Returns (morse_str, decoded_text).
    """
    on_ms = sorted([ms for v, ms in runs_ms if v == 1])
    if not on_ms:
        return '', ''

    # Find dot/dash boundary = biggest gap in sorted ON durations
    def split_two(vals):
        if len(vals) < 2:
            return vals, []
        gaps = [(vals[i+1] - vals[i], i) for i in range(len(vals)-1)]
        idx = max(gaps)[1] + 1
        return vals[:idx], vals[idx:]

    dots_ms, dashes_ms = split_two(on_ms)
    dot_unit = float(np.mean(dots_ms)) if dots_ms else float(np.mean(on_ms)) / 3
    dot_thresh  = dot_unit * 2.0   # ON  > 2x dot = dash
    intra_thresh = dot_unit * 2.0  # OFF < 2x dot = intra-symbol gap
    char_thresh  = dot_unit * 5.0  # OFF 2–5x dot = char gap; above = word gap

    morse_words, current_word, current_char = [], [], []

    def flush_char():
        if current_char:
            current_word.append(''.join(current_char))
            current_char.clear()

    def flush_word():
        flush_char()
        if current_word:
            morse_words.append(list(current_word))
            current_word.clear()

    for val, ms in runs_ms:
        if val == 1:
            current_char.append('.' if ms <= dot_thresh else '-')
        else:
            if ms <= intra_thresh:
                pass
            elif ms <= char_thresh:
                flush_char()
            else:
                flush_word()

    flush_word()

    morse_str = ' / '.join(' '.join(w) for w in morse_words)
    decoded   = ' '.join(
        ''.join(REVERSE_MORSE.get(ch, f'[{ch}]') for ch in word)
        for word in morse_words
    )
    return morse_str, decoded
def _is_garbage(decoded: str) -> bool:
    """Return True if a decoded result looks like garbage output."""
    if not decoded.strip():
        return True
    chars = decoded.replace(' ', '')
    if not chars:
        return True
    # If over 60% of characters are '?' it's garbage
    if chars.count('?') / len(chars) > 0.6:
        return True
    # If the only unique chars are E, T, I, A, N (Morse tree artefacts)
    unique = set(chars.upper())
    tree_chars = {'E', 'T', 'I', 'A', 'N', 'M', 'S', 'U', 'R', 'W', 'D', 'K'}
    if unique.issubset(tree_chars) and len(unique) <= 5:
        return True
    return False


def _decode_signal_processing(wav_path: str):
    """Custom signal-processing decoder with auto tone-frequency detection."""
    sr, samples = _read_wav_samples(wav_path)

    # 1. Auto-detect the carrier tone frequency
    lo, hi, peak_freq = _detect_tone_frequency(samples, sr)

    # 2. Compute band-limited energy (5ms hop)
    hop_ms = 5
    energy = _bandpass_energy(samples, sr, lo=lo, hi=hi, window_ms=20)
    energy = _smooth(energy, k=5)

    # 3. Binarise ON/OFF
    binary = _energy_to_binary(energy)

    # 4. Get run lengths in milliseconds
    runs_ms = _run_lengths_ms(binary, hop_ms=hop_ms)

    # 5. Filter noise blips < 15ms
    runs_ms = [(v, ms) for v, ms in runs_ms if ms >= 15]

    if not any(v == 1 for v, _ in runs_ms):
        return None, None, None

    # 6. Auto-classify and decode
    morse_str, decoded = _auto_classify_and_decode(runs_ms)
    return morse_str, decoded, f'Signal Processing (tone: {peak_freq:.0f} Hz)'


def _decode_library(wav_path: str):
    """morse_audio_decoder library decoder. Returns (morse_str, decoded, method_label)."""
    if not LIB_DECODER_AVAILABLE:
        return None, None, None
    try:
        mc = _LibMorseCode.from_wavfile(wav_path)
        on_samples, off_samples = mc._on_off_samples()
        dash_dot_chars           = mc._dash_dot_characters(on_samples)
        char_break_idx, word_space_idx = mc._break_spaces(off_samples)
        morse_words              = mc._morse_words(dash_dot_chars, char_break_idx, word_space_idx)
        morse_str = ' / '.join(' '.join(w) for w in morse_words)
        decoded   = mc._translate(morse_words)
        return morse_str, decoded, 'Library Decoder'
    except Exception:
        return None, None, None


def decode_audio_file(path: str):
    """
    Try both decoders. Use whichever gives a clean result.
    Returns (morse_str, decoded_text, method_label).
    """
    wav_path, is_temp = load_audio_as_mono_wav(path)
    try:
        results = []

        # 1. Try custom signal processing
        ms1, dec1, lbl1 = _decode_signal_processing(wav_path)
        if ms1 is not None:
            results.append((ms1, dec1, lbl1))

        # 2. Try library decoder (needs its own mono wav at 44100 for best results)
        if LIB_DECODER_AVAILABLE:
            try:
                if PYDUB_AVAILABLE:
                    audio = AudioSegment.from_wav(wav_path)
                    audio = audio.set_channels(1).set_frame_rate(44100)
                    tmp2 = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
                    audio.export(tmp2.name, format='wav')
                    ms2, dec2, lbl2 = _decode_library(tmp2.name)
                    os.unlink(tmp2.name)
                else:
                    ms2, dec2, lbl2 = _decode_library(wav_path)
                if ms2 is not None:
                    results.append((ms2, dec2, lbl2))
            except Exception:
                pass

        if not results:
            return '', '(No Morse signal detected — check your audio file)', 'None'

        # Pick the best result: prefer whichever has fewer '?' characters
        def score(r):
            txt = r[1].replace(' ', '')
            if not txt:
                return 0
            return 1 - (txt.count('?') / len(txt))

        best = max(results, key=score)
        morse_str, decoded, method = best

        # If both produced non-garbage and they differ, show both
        good_results = [r for r in results if not _is_garbage(r[1])]
        if len(good_results) == 2 and good_results[0][1] != good_results[1][1]:
            combined_decoded = (
                f"[{good_results[0][2]}]\n{good_results[0][1]}\n\n"
                f"[{good_results[1][2]}]\n{good_results[1][1]}"
            )
            return good_results[0][0], combined_decoded, 'Both'

        if _is_garbage(decoded):
            decoded = ('(Could not decode — audio may be too noisy '
                       'or not contain clean Morse code tones.\n'
                       'Both decoders were tried.)')

        return morse_str, decoded, method

    finally:
        if is_temp:
            os.unlink(wav_path)


def morse_to_text(morse_code: str) -> str:
    """Convert a manually-typed Morse code string to plain text."""
    words = morse_code.strip().split(' / ')
    decoded_words = []
    for word in words:
        letters = word.strip().split(' ')
        decoded_word = ''.join(REVERSE_MORSE.get(ch, '?') for ch in letters if ch)
        decoded_words.append(decoded_word)
    return ' '.join(decoded_words)


def save_to_word(morse_code: str, text: str, filename: str):
    doc = docx.Document()
    doc.add_heading('Morse Code Decoder Output', 0)
    doc.add_heading('Morse Code', level=1)
    doc.add_paragraph(morse_code)
    doc.add_heading('Decoded Text', level=1)
    doc.add_paragraph(text)
    doc.save(filename)

# ─── GUI ──────────────────────────────────────────────────────────────────────

class MorseDecoderApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('Morse Code Audio Decoder')
        self.geometry('720x620')
        self.resizable(True, True)
        self.configure(bg='#1e1e2e')

        self._dropped_file = None
        self._build_ui()
        self._setup_drag_and_drop()

    # ── UI construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        ACCENT  = '#89b4fa'
        BG      = '#1e1e2e'
        SURFACE = '#313244'
        TEXT    = '#cdd6f4'
        SUBTEXT = '#a6adc8'

        title_font = tkFont.Font(family='Segoe UI', size=16, weight='bold')
        body_font  = tkFont.Font(family='Segoe UI', size=10)
        mono_font  = tkFont.Font(family='Consolas',  size=10)

        # Title
        tk.Label(self, text='🎙 Morse Code Audio Decoder', font=title_font,
                 bg=BG, fg=ACCENT).pack(pady=(20, 4))
        tk.Label(self,
                 text='Drag & drop an audio file below, or click "Browse"',
                 font=body_font, bg=BG, fg=SUBTEXT).pack()

        ext_list = '  '.join(SUPPORTED_EXTENSIONS)
        tk.Label(self, text=f'Supported formats: {ext_list}',
                 font=body_font, bg=BG, fg=SUBTEXT).pack(pady=(0, 12))

        # Drop zone – Canvas gives us a true dashed border
        self._drop_canvas = tk.Canvas(self, bg=SURFACE, highlightthickness=0,
                                      height=90, cursor='hand2')
        self._drop_canvas.pack(fill='x', padx=30, pady=4)
        self._drop_canvas.bind('<Configure>', self._redraw_drop_border)

        # Inner frame sits inside the canvas
        self.drop_frame = tk.Frame(self._drop_canvas, bg=SURFACE, cursor='hand2')
        self._drop_canvas_window = self._drop_canvas.create_window(
            0, 0, anchor='nw', window=self.drop_frame
        )

        self.drop_label = tk.Label(
            self.drop_frame,
            text='⬇  Drop audio file here',
            font=tkFont.Font(family='Segoe UI', size=13),
            bg=SURFACE, fg=SUBTEXT
        )
        self.drop_label.pack(pady=(12, 2))

        self.file_label = tk.Label(self.drop_frame, text='',
                                   font=body_font, bg=SURFACE, fg=ACCENT,
                                   wraplength=580)
        self.file_label.pack(pady=(0, 10))

        # Buttons row
        btn_frame = tk.Frame(self, bg=BG)
        btn_frame.pack(pady=10)

        btn_style = dict(font=body_font, bg=ACCENT, fg='#1e1e2e',
                         relief='flat', padx=14, pady=6, cursor='hand2')

        tk.Button(btn_frame, text='Browse…',
                  command=self._browse, **btn_style).pack(side='left', padx=6)
        tk.Button(btn_frame, text='Decode Audio',
                  command=self._start_decode, **btn_style).pack(side='left', padx=6)
        tk.Button(btn_frame, text='Save as .docx',
                  command=self._save_docx, **btn_style).pack(side='left', padx=6)
        tk.Button(btn_frame, text='Clear',
                  command=self._clear, **btn_style).pack(side='left', padx=6)

        # Status bar
        self.status_var = tk.StringVar(value='Ready')
        tk.Label(self, textvariable=self.status_var,
                 font=body_font, bg=BG, fg=SUBTEXT).pack()

        # Progress bar
        self.progress = ttk.Progressbar(self, mode='indeterminate', length=300)
        self.progress.pack(pady=4)

        # Output areas
        def make_box(label_text):
            tk.Label(self, text=label_text, font=body_font,
                     bg=BG, fg=TEXT, anchor='w').pack(fill='x', padx=30, pady=(8,0))
            box = scrolledtext.ScrolledText(
                self, height=5, font=mono_font,
                bg=SURFACE, fg=TEXT, insertbackground=TEXT,
                relief='flat', bd=0, wrap='word'
            )
            box.pack(fill='both', padx=30, pady=(2,0), expand=True)
            return box

        self.morse_box  = make_box('Morse Code:')
        self.decode_box = make_box('Decoded Text:')

        # Store colours for drag feedback
        self._SURFACE = SURFACE
        self._ACCENT  = ACCENT

    # ── Drag-and-drop ────────────────────────────────────────────────────────

    def _setup_drag_and_drop(self):
        """Wire up drag-and-drop using TkinterDnD2 if available,
           otherwise fall back to click-to-browse on the drop zone."""
        try:
            self._drop_canvas.drop_target_register('DND_Files')
            self._drop_canvas.dnd_bind('<<Drop>>', self._on_drop)
            self._drop_canvas.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self._drop_canvas.dnd_bind('<<DragLeave>>', self._on_drag_leave)
        except Exception:
            self._drop_canvas.bind('<Button-1>', lambda e: self._browse())
            self.drop_frame.bind('<Button-1>', lambda e: self._browse())
            self.drop_label.bind('<Button-1>', lambda e: self._browse())
            self.drop_label.config(
                text='⬇  Click here to choose a file\n(install tkinterdnd2 for drag-and-drop)'
            )

    def _redraw_drop_border(self, event=None, colour=None):
        """Draw a dashed rectangle around the drop zone canvas."""
        c = self._drop_canvas
        w = c.winfo_width()
        h = c.winfo_height()
        # Resize the inner frame window to fill the canvas
        c.itemconfig(self._drop_canvas_window, width=w, height=h)
        c.delete('border')
        colour = colour or '#585b70'
        c.create_rectangle(4, 4, w - 4, h - 4,
                            outline=colour, width=2,
                            dash=(8, 4), tags='border')

    def _on_drag_enter(self, event):
        self._drop_canvas.config(bg=self._ACCENT)
        self.drop_frame.config(bg=self._ACCENT)
        self.drop_label.config(bg=self._ACCENT, fg='#1e1e2e')
        self._redraw_drop_border(colour=self._ACCENT)

    def _on_drag_leave(self, event):
        self._drop_canvas.config(bg=self._SURFACE)
        self.drop_frame.config(bg=self._SURFACE)
        self.drop_label.config(bg=self._SURFACE, fg='#a6adc8')
        self._redraw_drop_border(colour='#585b70')

    def _on_drop(self, event):
        self._on_drag_leave(event)
        path = event.data.strip().strip('{}')   # handle spaces in path
        self._set_file(path)

    # ── File handling ────────────────────────────────────────────────────────

    def _browse(self):
        from tkinter import filedialog
        ext_pairs = [('Audio files',
                      ' '.join(f'*{e}' for e in SUPPORTED_EXTENSIONS)),
                     ('All files', '*.*')]
        path = filedialog.askopenfilename(filetypes=ext_pairs)
        if path:
            self._set_file(path)

    def _set_file(self, path: str):
        ext = os.path.splitext(path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            messagebox.showerror(
                'Unsupported format',
                f'"{ext}" is not supported.\nPlease use: {", ".join(SUPPORTED_EXTENSIONS)}'
            )
            return
        self._dropped_file = path
        self.file_label.config(text=f'📄 {os.path.basename(path)}')
        self.status_var.set(f'File loaded: {os.path.basename(path)}')

    # ── Decode ───────────────────────────────────────────────────────────────

    def _start_decode(self):
        if not self._dropped_file:
            messagebox.showwarning('No file', 'Please load an audio file first.')
            return
        self.progress.start(10)
        self.status_var.set('Decoding…')
        threading.Thread(target=self._decode_worker, daemon=True).start()

    def _decode_worker(self):
        try:
            morse_str, decoded, method = decode_audio_file(self._dropped_file)
            self.after(0, self._show_results, morse_str, decoded, method)
        except Exception as exc:
            self.after(0, self._show_error, str(exc))

    def _show_results(self, morse_str, decoded, method):
        self.progress.stop()
        self._write(self.morse_box,  morse_str)
        self._write(self.decode_box, decoded)
        self.status_var.set(f'Decoding complete ✓  —  method: {method}')

    def _show_error(self, msg):
        self.progress.stop()
        self.status_var.set('Error – see dialog')
        messagebox.showerror('Decoding error', msg)

    # ── Save ─────────────────────────────────────────────────────────────────

    def _save_docx(self):
        morse_str = self.morse_box.get('1.0', 'end').strip()
        decoded   = self.decode_box.get('1.0', 'end').strip()
        if not morse_str and not decoded:
            messagebox.showwarning('Nothing to save', 'Decode an audio file first.')
            return
        from tkinter import filedialog
        path = filedialog.asksaveasfilename(
            defaultextension='.docx',
            filetypes=[('Word document', '*.docx')]
        )
        if path:
            save_to_word(morse_str, decoded, path)
            self.status_var.set(f'Saved → {os.path.basename(path)}')
            messagebox.showinfo('Saved', f'Document saved to:\n{path}')

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _clear(self):
        self._dropped_file = None
        self.file_label.config(text='')
        self._write(self.morse_box,  '')
        self._write(self.decode_box, '')
        self.status_var.set('Ready')

    @staticmethod
    def _write(widget, text):
        widget.config(state='normal')
        widget.delete('1.0', 'end')
        widget.insert('end', text)


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    # Try to enable native drag-and-drop via TkinterDnD2
    try:
        from tkinterdnd2 import TkinterDnD
        class MorseDecoderAppDnD(MorseDecoderApp, TkinterDnD.Tk):  # type: ignore
            def __init__(self):
                TkinterDnD.Tk.__init__(self)
                MorseDecoderApp._build_ui(self)
                MorseDecoderApp._setup_drag_and_drop(self)
                self.title('Morse Code Audio Decoder')
                self.geometry('720x620')
                self.configure(bg='#1e1e2e')
                self._dropped_file = None
        app = MorseDecoderAppDnD()
    except ImportError:
        app = MorseDecoderApp()

    app.mainloop()